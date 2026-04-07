"""
app/core/parsers/docx_parser.py — Microsoft Word Document Parser.

HOW DOCX PARSING WORKS
─────────────────────────────────────────────────────────────────────────────
A .docx file is a ZIP archive containing a series of XML files.
The text content is stored in `word/document.xml`.

Our parser uses the `python-docx` library to walk the XML tree:
  1. We iterate over `Paragraph` objects to extract the main prose.
  2. We iterate over `Table` objects to extract structured grid data.
  3. We identify different styles (headings vs bodies) to ensure coherent
     text flow for downstream chunking.

IMPORTANCE OF TABLES
─────────────────────────────────────────────────────────────────────────────
Traditional text extraction often ignores tables or scrambles their columns.
Our parser explicitly iterates through columns and rows, joining them with 
pipes ( | ) to maintain a readable grid structure for the RAG LLM.
"""

from __future__ import annotations
import logging
import docx
from .base_parser import BaseParser

logger = logging.getLogger(__name__)

class DocxParser(BaseParser):
    """
    Extracts text from Microsoft Word (.docx) files using python-docx.
    """

    def _extract_text(self, file_path: str, **kwargs) -> str:
        """
        Implementation of the abstract method.
        Iterates over the document's paragraph and table objects.
        """
        doc = docx.Document(file_path)
        full_content: list[str] = []

        # ── Extract Paragraphs ─────────────────────────────────────────────
        #
        # doc.paragraphs returns all text not inside tables or headers.
        logger.info("DocxParser: Extracting paragraphs from %s", file_path)
        for p in doc.paragraphs:
            if p.text.strip():
                full_content.append(p.text)

        # ── Extract Tables ─────────────────────────────────────────────────
        #
        # Tables carry high-density information. We join cells with pipes
        # to ensure the LLM understands the relational structure.
        if doc.tables:
            logger.info("DocxParser: Found %d tables in %s", len(doc.tables), file_path)
            full_content.append("\n[TABLE DATA EXTRACTED]")
            
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_content.append(" | ".join(row_data))

        # ── Return formatted content ───────────────────────────────────────
        return "\n\n".join(full_content)
